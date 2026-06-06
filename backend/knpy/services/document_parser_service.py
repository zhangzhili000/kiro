import os
import re
from typing import List, Dict, Optional
from fastapi import HTTPException


class DocumentParser:
    """文档解析服务，使用MinerU进行文档解析，支持多种格式"""
    
    def __init__(self):
        self.supported_formats = ["pdf", "docx", "doc", "md", "txt", "rtf", "pptx", "xlsx"]
    
    def parse_file(self, file_path: str, file_type: Optional[str] = None) -> Dict:
        """解析文件内容"""
        if not file_type:
            file_type = self._get_file_type(file_path)
        
        if file_type not in self.supported_formats:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")
        
        try:
            print(f"DEBUG: 尝试使用MinerU解析 {file_type} 文件")
            result = self._parse_with_mineru(file_path, file_type)
            print(f"DEBUG: MinerU解析完成，图片数: {len(result.get('images', []))}")
            return result
        except Exception as e:
            print(f"DEBUG: MinerU解析失败，降级到基础解析: {str(e)}")
            # 如果MinerU解析失败，降级到基础解析方法
            fallback_result = self._fallback_parse(file_path, file_type)
            if fallback_result:
                print(f"DEBUG: 基础解析完成，图片数: {len(fallback_result.get('images', []))}")
                return fallback_result
            raise HTTPException(status_code=500, detail=f"Failed to parse file: {str(e)}")
    
    def _get_file_type(self, file_path: str) -> str:
        """从文件路径获取文件类型"""
        ext = os.path.splitext(file_path)[1].lower().lstrip(".")
        return ext
    
    def _parse_with_mineru(self, file_path: str, file_type: str) -> Dict:
        """使用MinerU解析文件"""
        try:
            from mineru import DocumentParser as MinerUParser
            
            parser = MinerUParser()
            result = parser.parse(file_path)
            
            content = ""
            tables = []
            images = []
            
            # 提取文本内容（保留原始结构）
            if hasattr(result, 'content'):
                content = result.content
            elif hasattr(result, 'text'):
                content = result.text
            elif isinstance(result, dict):
                content = result.get('content', result.get('text', str(result)))
            else:
                content = str(result)
            
            # 提取表格（保留结构化数据）
            if hasattr(result, 'tables') and result.tables:
                for table in result.tables:
                    if isinstance(table, dict):
                        tables.append({
                            "content": table.get('content', ''),
                            "data": table.get('data', []),
                            "page": table.get('page', 0)
                        })
                    else:
                        tables.append({
                            "content": str(table),
                            "data": [],
                            "page": 0
                        })
            
            # 提取图片（保留base64或路径）
            if hasattr(result, 'images') and result.images:
                for img in result.images:
                    if isinstance(img, dict):
                        images.append({
                            "path": img.get('path', ''),
                            "base64": img.get('base64', ''),
                            "page": img.get('page', 0),
                            "width": img.get('width', 0),
                            "height": img.get('height', 0)
                        })
                    else:
                        images.append({
                            "path": str(img),
                            "base64": "",
                            "page": 0,
                            "width": 0,
                            "height": 0
                        })
            
            return {
                "content": content.strip(),
                "metadata": {
                    "tables_count": len(tables),
                    "images_count": len(images),
                    "parser": "mineru"
                },
                "tables": tables,
                "images": images,
                "type": file_type
            }
        
        except ImportError:
            raise RuntimeError("MinerU library not installed")
        except Exception as e:
            raise RuntimeError(f"MinerU parse error: {str(e)}")
    
    def _fallback_parse(self, file_path: str, file_type: str) -> Optional[Dict]:
        """降级解析方法，当MinerU不可用时使用"""
        try:
            if file_type == "pdf":
                return self._parse_pdf_fallback(file_path)
            elif file_type == "docx":
                return self._parse_docx_fallback(file_path)
            elif file_type == "md":
                return self._parse_md(file_path)
            elif file_type == "txt":
                return self._parse_txt(file_path)
            else:
                return None
        except Exception:
            return None
    
    def _parse_pdf_fallback(self, file_path: str) -> Dict:
        """PDF文件降级解析（使用PyMuPDF提取文字和图片）"""
        try:
            import fitz  # PyMuPDF
            import base64
            
            doc = fitz.open(file_path)
            content = ""
            images = []
            html_content = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # 提取文字
                text = page.get_text()
                if text:
                    content += text + "\n\n"
                
                # 提取页面中的图片对象（不是整个页面转图片）
                page_images = self._extract_pdf_images_with_positions(page, page_num)
                images.extend(page_images)
                
                # 生成带图片位置的HTML
                page_html = self._generate_pdf_page_html(page, page_images)
                html_content += page_html
            
            return {
                "content": content.strip(),
                "html_content": html_content,
                "metadata": {
                    "pages": len(doc),
                    "parser": "PyMuPDF"
                },
                "tables": [],
                "images": images,
                "type": "pdf"
            }
        except ImportError as e:
            print(f"DEBUG: PyMuPDF未安装: {str(e)}")
            # 降级到PyPDF2
            return self._parse_pdf_with_pypdf2(file_path)
        except Exception as e:
            print(f"DEBUG: PDF解析失败: {str(e)}")
            return self._parse_pdf_with_pypdf2(file_path)
    
    def _parse_pdf_with_pypdf2(self, file_path: str) -> Dict:
        """使用PyPDF2解析PDF（降级方案）"""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            content = ""
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content += text + "\n\n"
            
            return {
                "content": content.strip(),
                "metadata": {
                    "pages": len(reader.pages),
                    "parser": "PyPDF2"
                },
                "tables": [],
                "images": [],
                "type": "pdf"
            }
        except ImportError:
            return None
    
    def _extract_images_from_pdf_file(self, file_path: str) -> List[Dict]:
        """从PDF文件提取页面图像（使用PyMuPDF）"""
        images = []
        try:
            import fitz  # PyMuPDF
            import base64
            
            # 打开PDF文件
            doc = fitz.open(file_path)
            
            print(f"DEBUG: PDF文件打开成功，共 {len(doc)} 页")
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # 将页面渲染为图像
                pix = page.get_pixmap(dpi=150)
                
                # 转换为base64
                base64_data = base64.b64encode(pix.tobytes()).decode('utf-8')
                
                images.append({
                    "path": "",
                    "base64": base64_data,
                    "page": page_num + 1,
                    "width": pix.width,
                    "height": pix.height
                })
            
            print(f"DEBUG: PDF页面图像提取成功，共 {len(images)} 页")
            return images
        except ImportError as e:
            print(f"DEBUG: PyMuPDF未安装: {str(e)}")
            return []
        except Exception as e:
            print(f"DEBUG: PDF图像提取失败: {str(e)}")
            return []
    
    def _extract_pdf_images_with_positions(self, page, page_num: int) -> List[Dict]:
        """从PDF页面提取图片对象及其位置信息"""
        images = []
        try:
            import base64
            
            # 获取页面中的图片
            image_list = page.get_images(full=True)
            
            # 获取页面文本块信息（包含图片位置）
            blocks = page.get_text("blocks")
            
            # 建立图片索引到位置的映射
            img_positions = {}
            for block in blocks:
                # block格式: (x0, y0, x1, y1, "text", block_no, block_type)
                x0, y0, x1, y1, text, block_no, block_type = block
                if block_type == 1:  # 图像块
                    # 尝试从文本中提取图片索引信息
                    if text.startswith("img") or text.startswith("Image"):
                        try:
                            idx = int(text.replace("img", "").replace("Image", "").strip())
                            img_positions[idx] = (x0, y0, x1, y1)
                        except:
                            pass
            
            for img_index, img_info in enumerate(image_list):
                # img_info包含: (xref, smask, width, height, bpc, colorspace, alt. colorspace, name, filter, ...)
                xref = img_info[0]
                width = img_info[2]
                height = img_info[3]
                
                # 获取图片数据
                pix = page.parent.extract_image(xref)
                if pix and 'image' in pix:
                    base64_data = base64.b64encode(pix['image']).decode('utf-8')
                    
                    # 获取图片位置（从文本块映射或使用默认值）
                    x0, y0, x1, y1 = img_positions.get(img_index, (0, 0, width, height))
                    
                    images.append({
                        "path": "",
                        "base64": base64_data,
                        "page": page_num + 1,
                        "index": img_index,
                        "width": width,
                        "height": height,
                        "x0": x0,
                        "y0": y0,
                        "x1": x1,
                        "y1": y1,
                        "ext": pix.get('ext', 'png')
                    })
            
            print(f"DEBUG: 第 {page_num + 1} 页提取到 {len(images)} 张图片")
            return images
        except Exception as e:
            print(f"DEBUG: 提取图片失败: {str(e)}")
            return []
    
    def _generate_pdf_page_html(self, page, images: List[Dict]) -> str:
        """生成包含文字和图片的页面HTML"""
        elements = []
        
        # 获取页面文本块
        text_blocks = page.get_text("blocks")
        
        # 处理文本块
        for block in text_blocks:
            # block格式: (x0, y0, x1, y1, "text", block_no, block_type)
            x0, y0, x1, y1, text, block_no, block_type = block
            
            # 清理文本
            text = text.strip()
            if not text:
                continue
            
            # 根据块类型添加不同的HTML标签
            if block_type == 0:  # 文本
                # 尝试识别标题（字体大小较大）
                font_size = (y1 - y0)
                if font_size > 16:
                    content = f'<h2>{text}</h2>'
                elif font_size > 14:
                    content = f'<h3>{text}</h3>'
                else:
                    content = f'<p>{text}</p>'
                elements.append({
                    'type': 'text',
                    'y': y0,
                    'x': x0,
                    'content': content
                })
            elif block_type == 1:  # 图像占位
                # 将图片占位符添加到元素列表，后续会被实际图片替换
                elements.append({
                    'type': 'image_placeholder',
                    'y': y0,
                    'x': x0,
                    'block': block
                })
        
        # 添加图片到元素列表（使用图片自身的位置信息）
        for img in images:
            img_html = self._image_to_html(img)
            if img_html:
                # 使用图片在PDF中的实际位置
                y_pos = img.get('y0', 0)
                x_pos = img.get('x0', 0)
                elements.append({
                    'type': 'image',
                    'y': y_pos,
                    'x': x_pos,
                    'content': f'<div style="margin: 0.5em 0;">{img_html}</div>'
                })
        
        # 按位置排序（从上到下，从左到右）
        elements.sort(key=lambda e: (e['y'], e['x']))
        
        # 生成最终HTML
        html_parts = []
        for elem in elements:
            if elem['type'] != 'image_placeholder':  # 跳过图片占位符，使用实际图片
                html_parts.append(elem['content'])
        
        # 添加分页标记
        html_parts.append('<hr style="border: none; border-top: 1px dashed #ccc; margin: 2em 0;" />')
        
        return ''.join(html_parts)
    
    def _parse_docx_fallback(self, file_path: str) -> Dict:
        """DOCX文件降级解析"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n\n"
            
            return {
                "content": content.strip(),
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "parser": "python-docx"
                },
                "tables": [],
                "images": [],
                "type": "docx"
            }
        except ImportError:
            return None
    
    def _parse_md(self, file_path: str) -> Dict:
        """解析Markdown文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        headings = self._extract_headings(content)
        
        return {
            "content": content,
            "metadata": {
                "headings": headings,
                "parser": "native"
            },
            "tables": [],
            "images": [],
            "type": "md"
        }
    
    def _parse_txt(self, file_path: str) -> Dict:
        """解析TXT文件"""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        
        return {
            "content": content,
            "metadata": {
                "parser": "native"
            },
            "tables": [],
            "images": [],
            "type": "txt"
        }
    
    def _extract_headings(self, content: str) -> List[Dict]:
        """提取Markdown标题"""
        pattern = r"^(#{1,6})\s+(.+)$"
        headings = []
        
        for line in content.split("\n"):
            match = re.match(pattern, line)
            if match:
                level = len(match.group(1))
                text = match.group(2)
                headings.append({
                    "level": level,
                    "text": text
                })
        
        return headings
    
    def convert_to_html(self, content: str, file_type: str, tables: List = None, images: List = None) -> str:
        """将文档内容转换为HTML（支持表格和图片）"""
        if file_type == "md":
            return self._md_to_html(content)
        
        # 如果是PDF且有页面图像，优先显示图像
        if file_type == "pdf" and images and len(images) > 0:
            return self._pdf_with_images_to_html(content, images)
        
        html = self._plain_text_to_html(content)
        
        # 添加表格
        if tables:
            for i, table in enumerate(tables):
                table_html = self._table_to_html(table)
                if table_html:
                    html = html.replace(f"[TABLE_{i}]", table_html)
        
        # 添加图片
        if images:
            for i, img in enumerate(images):
                img_html = self._image_to_html(img)
                if img_html:
                    html = html.replace(f"[IMAGE_{i}]", img_html)
        
        return html
    
    def _pdf_with_images_to_html(self, content: str, images: List) -> str:
        """将PDF转换为包含页面图像的HTML"""
        html_parts = []
        
        # 添加文本内容作为参考
        if content.strip():
            html_parts.append('<div class="pdf-text-content" style="display: none;">')
            html_parts.append(f'<p>{content}</p>')
            html_parts.append('</div>')
        
        # 添加页面图像
        for img in images:
            img_html = self._image_to_html(img)
            if img_html:
                html_parts.append(f'<div class="pdf-page-image" style="margin: 1em 0; text-align: center;">')
                html_parts.append(img_html)
                if img.get('page'):
                    html_parts.append(f'<p style="text-align: center; color: #999; font-size: 12px;">第 {img["page"]} 页</p>')
                html_parts.append('</div>')
        
        return f'<div class="document-content pdf-document">{''.join(html_parts)}</div>'
    
    def _md_to_html(self, content: str) -> str:
        """Markdown转HTML"""
        try:
            import markdown
            return markdown.markdown(content, extensions=['extra', 'toc'])
        except ImportError:
            return self._plain_text_to_html(content)
    
    def _plain_text_to_html(self, content: str) -> str:
        """纯文本转HTML"""
        paragraphs = content.split("\n\n")
        html = ""
        
        for para in paragraphs:
            para = para.strip()
            if para:
                # 尝试识别表格结构
                table_html = self._detect_and_convert_table(para)
                if table_html:
                    html += table_html + "\n"
                else:
                    html += f"<p>{para}</p>\n"
        
        return f"<div class='document-content'>{html}</div>"
    
    def _detect_and_convert_table(self, text: str) -> str:
        """检测并转换文本中的表格结构"""
        lines = text.strip().split("\n")
        if len(lines) < 2:
            return ""
        
        # 检查是否有表格分隔行（包含多个-和|）
        has_separator = False
        for line in lines:
            if "|" in line and "-" in line:
                has_separator = True
                break
        
        if not has_separator:
            return ""
        
        # 尝试转换为HTML表格
        try:
            html = "<table border='1' cellpadding='4' cellspacing='0' style='border-collapse: collapse; margin: 1em 0;'>\n"
            
            for line in lines:
                if "|" in line:
                    cells = line.split("|")
                    cells = [c.strip() for c in cells if c.strip()]
                    
                    # 判断是否为分隔行
                    is_separator = all(c.replace("-", "").replace("=", "").replace(" ", "") == "" for c in cells)
                    
                    if is_separator:
                        continue
                    
                    html += "  <tr>\n"
                    for cell in cells:
                        html += f"    <td>{cell}</td>\n"
                    html += "  </tr>\n"
            
            html += "</table>"
            return html
        except:
            return ""
    
    def _table_to_html(self, table: Dict) -> str:
        """将表格数据转换为HTML"""
        if not table:
            return ""
        
        # 优先使用结构化数据
        data = table.get('data', [])
        content = table.get('content', '')
        
        if data and isinstance(data, list) and len(data) > 0:
            html = "<table border='1' cellpadding='4' cellspacing='0' style='border-collapse: collapse; margin: 1em 0;'>\n"
            for row in data:
                html += "  <tr>\n"
                if isinstance(row, list):
                    for cell in row:
                        html += f"    <td>{cell}</td>\n"
                else:
                    html += f"    <td>{row}</td>\n"
                html += "  </tr>\n"
            html += "</table>"
            return html
        elif content:
            # 尝试从文本内容解析表格
            return self._detect_and_convert_table(content)
        
        return ""
    
    def _image_to_html(self, image: Dict) -> str:
        """将图片转换为HTML img标签"""
        if not image:
            return ""
        
        base64_data = image.get('base64', '')
        path = image.get('path', '')
        width = image.get('width', 0)
        height = image.get('height', 0)
        
        style = "max-width: 100%; height: auto;"
        if width:
            style += f" max-width: {width}px;"
        
        if base64_data:
            return f'<img src="data:image/png;base64,{base64_data}" style="{style}" />'
        elif path:
            # 如果是本地路径，尝试读取并转换为base64
            try:
                if os.path.exists(path):
                    with open(path, 'rb') as f:
                        base64_data = f.read().encode('base64').decode()
                    return f'<img src="data:image/png;base64,{base64_data}" style="{style}" />'
            except:
                pass
        
        return f'<img src="{path}" style="{style}" alt="图片" />'
    
    def split_content(self, content: str, max_chunk_size: int = 500) -> List[str]:
        """将内容分块处理"""
        chunks = []
        paragraphs = content.split("\n\n")
        
        current_chunk = ""
        for para in paragraphs:
            if len(current_chunk) + len(para) < max_chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks


document_parser = DocumentParser()
