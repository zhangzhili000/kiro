import os
import re
import glob
import time
import requests
import tempfile
import zipfile
import json
from typing import List, Dict, Optional
from fastapi import HTTPException

class DocumentParser:
    """文档解析服务，使用MinerU Open API进行文档解析，支持多种格式"""
    
    # MinerU Open API 配置
    MINERU_API_BASE = "https://mineru.net/api/v4"
    MINERU_AGENT_API_BASE = "https://mineru.net/api/v1/agent"
    
    def __init__(self, db=None):
        self.supported_formats = ["pdf", "docx", "doc", "md", "txt", "rtf", "pptx", "xlsx"]
        self.db = db
        self._document_model_config = None
        # 默认API地址（备用）
        self.default_mineru_api_base = "https://mineru.net/api/v4"
        self.default_mineru_agent_api_base = "https://mineru.net/api/v1/agent"
    
    def set_db(self, db):
        """设置数据库连接"""
        self.db = db
    
    def _get_document_model_config(self) -> Optional[Dict]:
        """获取文档解析模型配置"""
        if self._document_model_config:
            return self._document_model_config
        
        if not self.db:
            return None
        
        try:
            from kiro.models.ai_models import ModelConfig
            config = self.db.query(ModelConfig).filter(
                ModelConfig.type == "document",
                ModelConfig.status == "active"
            ).first()
            
            if config:
                self._document_model_config = {
                    "api_type": config.api_type,
                    "model_id": config.model_id,
                    "api_key": config.api_key,
                    "api_base": config.api_base,
                    "description": config.description
                }
                return self._document_model_config
        except Exception as e:
            print(f"DEBUG: 获取文档解析模型配置失败: {str(e)}")
        
        return None
    
    def parse_file(self, file_path: str, file_type: Optional[str] = None) -> Dict:
        """解析文件内容"""
        if not file_type:
            file_type = self._get_file_type(file_path)
        
        if file_type not in self.supported_formats:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")
        
        try:
            print(f"DEBUG: 尝试使用MinerU解析 {file_type} 文件")
            result = self._parse_with_mineru(file_path, file_type)
            print(f"DEBUG: MinerU解析完成，图片数: {len(result.get('images', []))}")
            return result
        except Exception as e:
            print(f"DEBUG: MinerU解析失败，降级到基础解析: {str(e)}")
            # 如果MinerU解析失败，降级到基础解析方法
            fallback_result = self._fallback_parse(file_path, file_type)
            if fallback_result:
                print(f"DEBUG: 基础解析完成，图片数: {len(fallback_result.get('images', []))}")
                return fallback_result
            raise HTTPException(status_code=500, detail=f"Failed to parse file: {str(e)}")
    
    def _get_file_type(self, file_path: str) -> str:
        """从文件路径获取文件类型"""
        ext = os.path.splitext(file_path)[1].lower().lstrip(".")
        return ext
    
    def _parse_with_mineru(self, file_path: str, file_type: str) -> Dict:
        """使用MinerU Open API解析文件"""
        model_config = self._get_document_model_config()
        
        # 检查是否配置了MinerU（api_type为mineru且存在）
        is_mineru_configured = model_config and model_config.get("api_type") == "mineru"
        
        if is_mineru_configured:
            token = model_config.get("api_key")
            api_base = model_config.get("api_base")
            parse_mode = model_config.get("model_id")  # agent 或 precision
            
            if parse_mode == "precision" or (parse_mode is None and token):
                # 用户选择精准解析，或者没有选择模式但配置了Token
                print(f"DEBUG: 使用MinerU精准解析API（有Token）")
                return self._parse_with_mineru_precision(file_path, file_type, token, api_base)
            else:
                # 用户选择轻量解析，或者没有选择模式且没有Token
                print(f"DEBUG: 使用MinerU轻量解析API（无Token）")
                return self._parse_with_mineru_agent(file_path, file_type, api_base)
        else:
            # 没有配置MinerU或配置为空，使用默认的轻量解析API
            print(f"DEBUG: 使用MinerU轻量解析API（默认方案）")
            return self._parse_with_mineru_agent(file_path, file_type)
    
    def _parse_with_mineru_precision(self, file_path: str, file_type: str, token: str, api_base: str = None) -> Dict:
        """使用MinerU精准解析API（需要Token）"""
        try:
            # 使用配置的API地址或默认地址
            base_url = api_base if api_base else self.default_mineru_api_base
            
            # Step 1: 申请上传URL
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {token}"
            }
            
            file_name = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            
            # 申请批量上传URL
            apply_url = f"{base_url}/file-urls/batch"
            apply_data = {
                "file_names": [file_name],
                "enable_formula": True,
                "enable_table": True,
                "model_version": "vlm"  # 使用VLM模型，精度更高
            }
            
            response = requests.post(apply_url, headers=headers, json=apply_data, timeout=30)
            if response.status_code != 200:
                raise RuntimeError(f"MinerU申请上传URL失败: {response.text}")
            
            result = response.json()
            batch_id = result.get("data", {}).get("batch_id")
            upload_urls = result.get("data", {}).get("file_urls", [])
            
            if not batch_id or not upload_urls:
                raise RuntimeError(f"MinerU返回数据无效: {result}")
            
            upload_url = upload_urls[0]
            print(f"DEBUG: MinerU申请上传URL成功，batch_id: {batch_id}")
            
            # Step 2: 上传文件
            with open(file_path, "rb") as f:
                upload_response = requests.put(upload_url, data=f.read(), timeout=60)
                if upload_response.status_code != 200:
                    raise RuntimeError(f"MinerU上传文件失败: {upload_response.text}")
            
            print(f"DEBUG: MinerU上传文件成功")
            
            # Step 3: 轮询结果
            max_retries = 60  # 最大轮询次数
            retry_interval = 5  # 轮询间隔（秒）
            
            for i in range(max_retries):
                time.sleep(retry_interval)
                
                status_url = f"{base_url}/extract/task?batch_id={batch_id}"
                status_response = requests.get(status_url, headers=headers, timeout=30)
                
                if status_response.status_code != 200:
                    continue
                
                status_result = status_response.json()
                tasks = status_result.get("data", {}).get("tasks", [])
                
                if not tasks:
                    continue
                
                task = tasks[0]
                task_status = task.get("status")
                
                if task_status == "completed":
                    # 获取下载URL
                    full_zip_url = task.get("full_zip_url")
                    if full_zip_url:
                        print(f"DEBUG: MinerU解析完成，开始下载结果")
                        return self._download_and_parse_mineru_result(full_zip_url, file_type)
                    else:
                        raise RuntimeError("MinerU返回结果缺少下载URL")
                
                elif task_status == "failed":
                    error_msg = task.get("error", "未知错误")
                    raise RuntimeError(f"MinerU解析失败: {error_msg}")
                
                print(f"DEBUG: MinerU解析进行中，状态: {task_status}，轮询次数: {i+1}")
            
            raise RuntimeError("MinerU解析超时")
        
        except requests.exceptions.Timeout:
            raise RuntimeError("MinerU API请求超时")
        except Exception as e:
            raise RuntimeError(f"MinerU精准解析错误: {str(e)}")
    
    def _parse_with_mineru_agent(self, file_path: str, file_type: str, api_base: str = None) -> Dict:
        """使用MinerU轻量解析API（免登录，IP限频）"""
        try:
            # 使用配置的API地址或默认地址
            base_url = api_base if api_base else self.default_mineru_agent_api_base
            
            # Step 1: 获取上传URL和task_id
            upload_url = f"{base_url}/parse/file"
            
            file_name = os.path.basename(file_path)
            file_size = os.path.getsize(file_path)
            
            # 检查文件大小限制（轻量API限制10MB）
            if file_size > 10 * 1024 * 1024:
                raise RuntimeError("文件大小超过10MB，轻量解析API不支持，请配置MinerU Token使用精准解析API")
            
            # 使用JSON请求体格式获取上传URL
            headers = {
                "Content-Type": "application/json"
            }
            
            data = {
                "file_name": file_name,
                "file": ""
            }
            
            response = requests.post(upload_url, headers=headers, json=data, timeout=60)
            
            if response.status_code != 200:
                raise RuntimeError(f"MinerU轻量解析获取上传URL失败: {response.text}")
            
            result = response.json()
            task_id = result.get("data", {}).get("task_id")
            oss_upload_url = result.get("data", {}).get("file_url")
            
            if not task_id or not oss_upload_url:
                raise RuntimeError(f"MinerU返回数据无效: {result}")
            
            print(f"DEBUG: MinerU轻量解析获取上传URL成功，task_id: {task_id}")
            
            # Step 2: 使用返回的OSS URL上传文件
            with open(file_path, "rb") as f:
                file_content = f.read()
            
            upload_response = requests.put(oss_upload_url, data=file_content, timeout=60)
            
            if upload_response.status_code != 200:
                raise RuntimeError(f"MinerU轻量解析上传文件到OSS失败: {upload_response.status_code}")
            
            print(f"DEBUG: MinerU轻量解析上传文件到OSS成功")
            
            # Step 3: 轮询结果
            max_retries = 30  # 最大轮询次数
            retry_interval = 3  # 轮询间隔（秒）
            
            for i in range(max_retries):
                time.sleep(retry_interval)
                
                status_url = f"{base_url}/parse/{task_id}"
                status_response = requests.get(status_url, timeout=30)
                
                if status_response.status_code != 200:
                    continue
                
                status_result = status_response.json()
                task_data = status_result.get("data", {})
                task_status = task_data.get("state")
                
                if task_status == "completed" or task_status == "done":
                    # 获取Markdown内容
                    markdown_url = task_data.get("markdown_url")
                    if markdown_url:
                        print(f"DEBUG: MinerU轻量解析完成，开始下载Markdown")
                        
                        # 下载Markdown内容
                        md_response = requests.get(markdown_url, timeout=30)
                        if md_response.status_code == 200:
                            content = md_response.text
                            return {
                                "content": content.strip(),
                                "metadata": {
                                    "parser": "mineru_agent",
                                    "task_id": task_id
                                },
                                "tables": [],
                                "images": [],
                                "type": file_type
                            }
                        else:
                            raise RuntimeError("MinerU Markdown下载失败")
                    else:
                        raise RuntimeError("MinerU返回结果缺少Markdown URL")
                
                elif task_status == "failed":
                    # 获取错误信息，API可能返回err_msg或error
                    error_msg = task_data.get("err_msg") or task_data.get("error") or "未知错误"
                    
                    # 检查是否是页数超限错误
                    if "page count exceeds API limit" in error_msg or "page_range" in error_msg:
                        print(f"DEBUG: MinerU轻量解析页数超限，开始分割PDF: {error_msg}")
                        raise RuntimeError(f"page_limit_exceeded: {error_msg}")
                    
                    raise RuntimeError(f"MinerU轻量解析失败: {error_msg}")
                
                print(f"DEBUG: MinerU轻量解析进行中，状态: {task_status}，轮询次数: {i+1}")
            
            raise RuntimeError("MinerU轻量解析超时")
        
        except RuntimeError as e:
            error_str = str(e)
            # 检查是否是页数超限错误
            if "page_limit_exceeded" in error_str and file_type == "pdf":
                # 尝试分割PDF并重新解析
                split_files = self._split_pdf(file_path)
                if len(split_files) > 1:
                    print(f"DEBUG: PDF已分割为 {len(split_files)} 个文件，开始逐个解析")
                    results = []
                    for split_file in split_files:
                        try:
                            result = self._parse_with_mineru_agent(split_file, file_type, api_base)
                            results.append(result)
                        except Exception as split_e:
                            print(f"DEBUG: 分割文件 {split_file} 解析失败: {str(split_e)}")
                            # 如果某个分割文件解析失败，尝试使用降级解析
                            fallback_result = self._fallback_parse(split_file, file_type)
                            if fallback_result:
                                results.append(fallback_result)
                    
                    if results:
                        merged_result = self._merge_parse_results(results)
                        print(f"DEBUG: PDF分割解析完成，已合并 {len(results)} 个结果")
                        return merged_result
            
            # 非页数超限错误，重新抛出
            raise e
        
        except requests.exceptions.Timeout:
            raise RuntimeError("MinerU API请求超时")
        except Exception as e:
            raise RuntimeError(f"MinerU轻量解析错误: {str(e)}")
    
    def _split_pdf(self, file_path: str, max_pages: int = 20) -> List[str]:
        """将PDF分割成多个小文件"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            total_pages = len(doc)
            split_files = []
            
            if total_pages <= max_pages:
                return [file_path]
            
            print(f"DEBUG: PDF文件共 {total_pages} 页，超过限制 {max_pages} 页，开始分割")
            
            for i in range(0, total_pages, max_pages):
                start_page = i
                end_page = min(i + max_pages, total_pages)
                
                output_path = os.path.join(
                    tempfile.gettempdir(),
                    f"split_{os.path.basename(file_path)}_{start_page + 1}_{end_page}.pdf"
                )
                
                new_doc = fitz.open()
                new_doc.insert_pdf(doc, from_page=start_page, to_page=end_page - 1)
                new_doc.save(output_path)
                new_doc.close()
                
                split_files.append(output_path)
                print(f"DEBUG: 已分割PDF页面 {start_page + 1}-{end_page} 到 {output_path}")
            
            doc.close()
            return split_files
        
        except ImportError:
            print("DEBUG: PyMuPDF未安装，无法分割PDF")
            return [file_path]
        except Exception as e:
            print(f"DEBUG: PDF分割失败: {str(e)}")
            return [file_path]
    
    def _merge_parse_results(self, results: List[Dict]) -> Dict:
        """合并多个解析结果"""
        merged_content = ""
        merged_tables = []
        merged_images = []
        page_offset = 0
        
        for i, result in enumerate(results):
            content = result.get("content", "")
            if content:
                if i > 0:
                    merged_content += "\n\n---\n\n"
                merged_content += content
            
            for table in result.get("tables", []):
                table_copy = table.copy()
                if "page" in table_copy:
                    table_copy["page"] += page_offset
                merged_tables.append(table_copy)
            
            for image in result.get("images", []):
                image_copy = image.copy()
                if "page" in image_copy:
                    image_copy["page"] += page_offset
                merged_images.append(image_copy)
            
            pages = result.get("metadata", {}).get("pages", 0)
            page_offset += pages
        
        return {
            "content": merged_content.strip(),
            "metadata": {
                "parser": "mineru_split",
                "tables_count": len(merged_tables),
                "images_count": len(merged_images)
            },
            "tables": merged_tables,
            "images": merged_images,
            "type": "pdf"
        }
    
    def _download_and_parse_mineru_result(self, zip_url: str, file_type: str) -> Dict:
        """下载并解析MinerU结果ZIP包"""
        try:
            # 下载ZIP文件
            response = requests.get(zip_url, timeout=60)
            if response.status_code != 200:
                raise RuntimeError(f"MinerU结果下载失败: {response.status_code}")
            
            # 解压ZIP文件
            with tempfile.TemporaryDirectory() as tmp_dir:
                zip_path = os.path.join(tmp_dir, "result.zip")
                with open(zip_path, "wb") as f:
                    f.write(response.content)
                
                with zipfile.ZipFile(zip_path, "r") as zf:
                    zf.extractall(tmp_dir)
                
                # 查找Markdown文件
                content = ""
                md_files = glob.glob(os.path.join(tmp_dir, "**", "*.md"), recursive=True)
                for md_file in md_files:
                    if os.path.basename(md_file) == "full.md" or "markdown" in md_file:
                        with open(md_file, "r", encoding="utf-8") as f:
                            content = f.read()
                        break
                
                # 查找JSON文件（表格、图片等）
                tables = []
                images = []
                json_files = glob.glob(os.path.join(tmp_dir, "**", "*.json"), recursive=True)
                
                for json_file in json_files:
                    try:
                        with open(json_file, "r", encoding="utf-8") as f:
                            data = json.load(f)
                            
                            if isinstance(data, dict):
                                # 提取表格
                                if "tables" in data:
                                    for table in data["tables"]:
                                        tables.append({
                                            "content": table.get('content', ''),
                                            "data": table.get('data', []),
                                            "page": table.get('page', 0)
                                        })
                                
                                # 提取图片信息
                                if "images" in data:
                                    for img in data["images"]:
                                        images.append({
                                            "path": img.get('path', ''),
                                            "base64": img.get('base64', ''),
                                            "page": img.get('page', 0),
                                            "width": img.get('width', 0),
                                            "height": img.get('height', 0)
                                        })
                    except:
                        continue
                
                return {
                    "content": content.strip(),
                    "metadata": {
                        "tables_count": len(tables),
                        "images_count": len(images),
                        "parser": "mineru_precision"
                    },
                    "tables": tables,
                    "images": images,
                    "type": file_type
                }
        
        except Exception as e:
            raise RuntimeError(f"MinerU结果解析错误: {str(e)}")
    
    def _fallback_parse(self, file_path: str, file_type: str) -> Optional[Dict]:
        """降级解析方法，当MinerU不可用时使用"""
        try:
            if file_type == "pdf":
                return self._parse_pdf_fallback(file_path)
            elif file_type == "docx":
                return self._parse_docx_fallback(file_path)
            elif file_type == "md":
                return self._parse_md(file_path)
            elif file_type == "txt":
                return self._parse_txt(file_path)
            else:
                return None
        except Exception:
            return None
    
    def _parse_pdf_fallback(self, file_path: str) -> Dict:
        """PDF文件降级解析（使用PyMuPDF提取文字和图片）"""
        try:
            import fitz  # PyMuPDF
            import base64
            
            doc = fitz.open(file_path)
            content = ""
            images = []
            html_content = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # 提取文字
                text = page.get_text()
                if text:
                    content += text + "\n\n"
                
                # 提取页面中的图片对象（不是整个页面转图片）
                page_images = self._extract_pdf_images_with_positions(page, page_num)
                images.extend(page_images)
                
                # 生成带图片位置的HTML
                page_html = self._generate_pdf_page_html(page, page_images)
                html_content += page_html
            
            return {
                "content": content.strip(),
                "html_content": html_content,
                "metadata": {
                    "pages": len(doc),
                    "parser": "PyMuPDF"
                },
                "tables": [],
                "images": images,
                "type": "pdf"
            }
        except ImportError as e:
            print(f"DEBUG: PyMuPDF未安装: {str(e)}")
            # 降级到PyPDF2
            return self._parse_pdf_with_pypdf2(file_path)
        except Exception as e:
            print(f"DEBUG: PDF解析失败: {str(e)}")
            return self._parse_pdf_with_pypdf2(file_path)
    
    def _parse_pdf_with_pypdf2(self, file_path: str) -> Dict:
        """使用PyPDF2解析PDF（降级方案）"""
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            content = ""
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    content += text + "\n\n"
            
            return {
                "content": content.strip(),
                "metadata": {
                    "pages": len(reader.pages),
                    "parser": "PyPDF2"
                },
                "tables": [],
                "images": [],
                "type": "pdf"
            }
        except ImportError:
            return None
    
    def _extract_images_from_pdf_file(self, file_path: str) -> List[Dict]:
        """从PDF文件提取页面图像（使用PyMuPDF）"""
        images = []
        try:
            import fitz  # PyMuPDF
            import base64
            
            # 打开PDF文件
            doc = fitz.open(file_path)
            
            print(f"DEBUG: PDF文件打开成功，共 {len(doc)} 页")
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # 将页面渲染为图像
                pix = page.get_pixmap(dpi=150)
                
                # 转换为base64
                base64_data = base64.b64encode(pix.tobytes()).decode('utf-8')
                
                images.append({
                    "path": "",
                    "base64": base64_data,
                    "page": page_num + 1,
                    "width": pix.width,
                    "height": pix.height
                })
            
            print(f"DEBUG: PDF页面图像提取成功，共 {len(images)} 页")
            return images
        except ImportError as e:
            print(f"DEBUG: PyMuPDF未安装: {str(e)}")
            return []
        except Exception as e:
            print(f"DEBUG: PDF图像提取失败: {str(e)}")
            return []
    
    def _extract_pdf_images_with_positions(self, page, page_num: int) -> List[Dict]:
        """从PDF页面提取图片对象及其位置信息"""
        images = []
        try:
            import base64
            
            # 获取页面中的图片
            image_list = page.get_images(full=True)
            
            # 获取页面文本块信息（包含图片位置）
            blocks = page.get_text("blocks")
            
            # 建立图片索引到位置的映射
            img_positions = {}
            for block in blocks:
                # block格式: (x0, y0, x1, y1, "text", block_no, block_type)
                x0, y0, x1, y1, text, block_no, block_type = block
                if block_type == 1:  # 图像块
                    # 尝试从文本中提取图片索引信息
                    if text.startswith("img") or text.startswith("Image"):
                        try:
                            idx = int(text.replace("img", "").replace("Image", "").strip())
                            img_positions[idx] = (x0, y0, x1, y1)
                        except:
                            pass
            
            for img_index, img_info in enumerate(image_list):
                # img_info包含: (xref, smask, width, height, bpc, colorspace, alt. colorspace, name, filter, ...)
                xref = img_info[0]
                width = img_info[2]
                height = img_info[3]
                
                # 获取图片数据
                pix = page.parent.extract_image(xref)
                if pix and 'image' in pix:
                    base64_data = base64.b64encode(pix['image']).decode('utf-8')
                    
                    # 获取图片位置（从文本块映射或使用默认值）
                    x0, y0, x1, y1 = img_positions.get(img_index, (0, 0, width, height))
                    
                    images.append({
                        "path": "",
                        "base64": base64_data,
                        "page": page_num + 1,
                        "index": img_index,
                        "width": width,
                        "height": height,
                        "x0": x0,
                        "y0": y0,
                        "x1": x1,
                        "y1": y1,
                        "ext": pix.get('ext', 'png')
                    })
            
            print(f"DEBUG: 第 {page_num + 1} 页提取到 {len(images)} 张图片")
            return images
        except Exception as e:
            print(f"DEBUG: 提取图片失败: {str(e)}")
            return []
    
    def _generate_pdf_page_html(self, page, images: List[Dict]) -> str:
        """生成包含文字和图片的页面HTML"""
        elements = []
        
        # 获取页面文本块
        text_blocks = page.get_text("blocks")
        
        # 处理文本块
        for block in text_blocks:
            # block格式: (x0, y0, x1, y1, "text", block_no, block_type)
            x0, y0, x1, y1, text, block_no, block_type = block
            
            # 清理文本
            text = text.strip()
            if not text:
                continue
            
            # 根据块类型添加不同的HTML标签
            if block_type == 0:  # 文本
                # 尝试识别标题（字体大小较大）
                font_size = (y1 - y0)
                if font_size > 16:
                    content = f'<h2>{text}</h2>'
                elif font_size > 14:
                    content = f'<h3>{text}</h3>'
                else:
                    content = f'<p>{text}</p>'
                elements.append({
                    'type': 'text',
                    'y': y0,
                    'x': x0,
                    'content': content
                })
            elif block_type == 1:  # 图像占位
                # 将图片占位符添加到元素列表，后续会被实际图片替换
                elements.append({
                    'type': 'image_placeholder',
                    'y': y0,
                    'x': x0,
                    'block': block
                })
        
        # 添加图片到元素列表（使用图片自身的位置信息）
        for img in images:
            img_html = self._image_to_html(img)
            if img_html:
                # 使用图片在PDF中的实际位置
                y_pos = img.get('y0', 0)
                x_pos = img.get('x0', 0)
                elements.append({
                    'type': 'image',
                    'y': y_pos,
                    'x': x_pos,
                    'content': f'<div style="margin: 0.5em 0;">{img_html}</div>'
                })
        
        # 按位置排序（从上到下，从左到右）
        elements.sort(key=lambda e: (e['y'], e['x']))
        
        # 生成最终HTML
        html_parts = []
        for elem in elements:
            if elem['type'] != 'image_placeholder':  # 跳过图片占位符，使用实际图片
                html_parts.append(elem['content'])
        
        # 添加分页标记
        html_parts.append('<hr style="border: none; border-top: 1px dashed #ccc; margin: 2em 0;" />')
        
        return ''.join(html_parts)
    
    def _parse_docx_fallback(self, file_path: str) -> Dict:
        """DOCX文件降级解析"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n\n"
            
            return {
                "content": content.strip(),
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "parser": "python-docx"
                },
                "tables": [],
                "images": [],
                "type": "docx"
            }
        except ImportError:
            return None
    
    def _parse_md(self, file_path: str) -> Dict:
        """解析Markdown文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        headings = self._extract_headings(content)
        
        return {
            "content": content,
            "metadata": {
                "headings": headings,
                "parser": "native"
            },
            "tables": [],
            "images": [],
            "type": "md"
        }
    
    def _parse_txt(self, file_path: str) -> Dict:
        """解析TXT文件"""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        
        return {
            "content": content,
            "metadata": {
                "parser": "native"
            },
            "tables": [],
            "images": [],
            "type": "txt"
        }
    
    def _extract_headings(self, content: str) -> List[Dict]:
        """提取Markdown标题"""
        pattern = r"^(#{1,6})\s+(.+)$"
        headings = []
        
        for line in content.split("\n"):
            match = re.match(pattern, line)
            if match:
                level = len(match.group(1))
                text = match.group(2)
                headings.append({
                    "level": level,
                    "text": text
                })
        
        return headings
    
    def convert_to_html(self, content: str, file_type: str, tables: List = None, images: List = None) -> str:
        """将文档内容转换为HTML（支持表格和图片）"""
        if file_type == "md":
            return self._md_to_html(content)
        
        # 如果是PDF且有页面图像，优先显示图像
        if file_type == "pdf" and images and len(images) > 0:
            return self._pdf_with_images_to_html(content, images)
        
        html = self._plain_text_to_html(content)
        
        # 添加表格
        if tables:
            for i, table in enumerate(tables):
                table_html = self._table_to_html(table)
                if table_html:
                    html = html.replace(f"[TABLE_{i}]", table_html)
        
        # 添加图片
        if images:
            for i, img in enumerate(images):
                img_html = self._image_to_html(img)
                if img_html:
                    html = html.replace(f"[IMAGE_{i}]", img_html)
        
        return html
    
    def _pdf_with_images_to_html(self, content: str, images: List) -> str:
        """将PDF转换为包含页面图像的HTML"""
        html_parts = []
        
        # 添加文本内容作为参考
        if content.strip():
            html_parts.append('<div class="pdf-text-content" style="display: none;">')
            html_parts.append(f'<p>{content}</p>')
            html_parts.append('</div>')
        
        # 添加页面图像
        for img in images:
            img_html = self._image_to_html(img)
            if img_html:
                html_parts.append(f'<div class="pdf-page-image" style="margin: 1em 0; text-align: center;">')
                html_parts.append(img_html)
                if img.get('page'):
                    html_parts.append(f'<p style="text-align: center; color: #999; font-size: 12px;">第 {img["page"]} 页</p>')
                html_parts.append('</div>')
        
        return f'<div class="document-content pdf-document">{"".join(html_parts)}</div>'
    
    def _md_to_html(self, content: str) -> str:
        """Markdown转HTML"""
        try:
            import markdown
            return markdown.markdown(content, extensions=['extra', 'toc'])
        except ImportError:
            return self._plain_text_to_html(content)
    
    def _plain_text_to_html(self, content: str) -> str:
        """纯文本转HTML"""
        paragraphs = content.split("\n\n")
        html = ""
        
        for para in paragraphs:
            para = para.strip()
            if para:
                # 尝试识别表格结构
                table_html = self._detect_and_convert_table(para)
                if table_html:
                    html += table_html + "\n"
                else:
                    html += f"<p>{para}</p>\n"
        
        return f"<div class='document-content'>{html}</div>"
    
    def _detect_and_convert_table(self, text: str) -> str:
        """检测并转换文本中的表格结构"""
        lines = text.strip().split("\n")
        if len(lines) < 2:
            return ""
        
        # 检查是否有表格分隔行（包含多个-和|）
        has_separator = False
        for line in lines:
            if "|" in line and "-" in line:
                has_separator = True
                break
        
        if not has_separator:
            return ""
        
        # 尝试转换为HTML表格
        try:
            html = "<table border='1' cellpadding='4' cellspacing='0' style='border-collapse: collapse; margin: 1em 0;'>\n"
            
            for line in lines:
                if "|" in line:
                    cells = line.split("|")
                    cells = [c.strip() for c in cells if c.strip()]
                    
                    # 判断是否为分隔行
                    is_separator = all(c.replace("-", "").replace("=", "").replace(" ", "") == "" for c in cells)
                    
                    if is_separator:
                        continue
                    
                    html += "  <tr>\n"
                    for cell in cells:
                        html += f"    <td>{cell}</td>\n"
                    html += "  </tr>\n"
            
            html += "</table>"
            return html
        except:
            return ""
    
    def _table_to_html(self, table: Dict) -> str:
        """将表格数据转换为HTML"""
        if not table:
            return ""
        
        # 优先使用结构化数据
        data = table.get('data', [])
        content = table.get('content', '')
        
        if data and isinstance(data, list) and len(data) > 0:
            html = "<table border='1' cellpadding='4' cellspacing='0' style='border-collapse: collapse; margin: 1em 0;'>\n"
            for row in data:
                html += "  <tr>\n"
                if isinstance(row, list):
                    for cell in row:
                        html += f"    <td>{cell}</td>\n"
                else:
                    html += f"    <td>{row}</td>\n"
                html += "  </tr>\n"
            html += "</table>"
            return html
        elif content:
            # 尝试从文本内容解析表格
            return self._detect_and_convert_table(content)
        
        return ""
    
    def _image_to_html(self, image: Dict) -> str:
        """将图片转换为HTML img标签"""
        if not image:
            return ""
        
        base64_data = image.get('base64', '')
        path = image.get('path', '')
        width = image.get('width', 0)
        height = image.get('height', 0)
        
        style = "max-width: 100%; height: auto;"
        if width:
            style += f" max-width: {width}px;"
        
        if base64_data:
            return f'<img src="data:image/png;base64,{base64_data}" style="{style}" />'
        elif path:
            # 如果是本地路径，尝试读取并转换为base64
            try:
                if os.path.exists(path):
                    with open(path, 'rb') as f:
                        base64_data = f.read().encode('base64').decode()
                    return f'<img src="data:image/png;base64,{base64_data}" style="{style}" />'
            except:
                pass
        
        return f'<img src="{path}" style="{style}" alt="图片" />'
    
    def split_content(self, content: str, max_chunk_size: int = 500) -> List[str]:
        """将内容分块处理"""
        chunks = []
        paragraphs = content.split("\n\n")
        
        current_chunk = ""
        for para in paragraphs:
            if len(current_chunk) + len(para) < max_chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks


document_parser = DocumentParser()
